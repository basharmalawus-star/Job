#!/usr/bin/env python3
"""
Job Application Assistant (MVP)
--------------------------------
Purpose:
  • Filter job lists by location/keywords
  • Tailor a resume & cover letter to each job description (rule‑based; no external APIs)
  • Open application links for semi‑automated submission (keeps you within site Terms of Service)

⚠️ Important
This script intentionally does NOT auto‑submit applications on platforms like LinkedIn or Handshake, which can violate their Terms of Service and get accounts flagged. Instead, it prepares tailored materials and opens the relevant links so you can quickly review and submit.

Quick Start
-----------
1) Create a Python virtual environment and install optional deps (only PyYAML needed):
    python -m venv .venv && . .venv/bin/activate   # Windows: .venv\\Scripts\\activate
    pip install pyyaml python-docx

2) Put your master resume data into master_resume.yaml (see SAMPLE at the bottom of this file).

3) Put your jobs into jobs.csv (see SAMPLE at the bottom of this file; you can export from trackers or hand‑add rows).

4) Run:
    python job_apply_mvp.py filter --locations "Sacramento, CA; Remote" --include "intern,analyst"
    python job_apply_mvp.py tailor --job-id 42 --out-dir out/
    python job_apply_mvp.py open --job-id 42   # opens apply_url in browser

Files produced in out/:
  • <jobid>_resume.md   – tailored markdown resume
  • <jobid>_resume.docx – tailored .docx (if python-docx installed)
  • <jobid>_cover_letter.md – tailored cover letter

"""

from __future__ import annotations
import argparse
import csv
import json
import os
import re
import sys
import textwrap
import webbrowser
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Iterable, Optional, Tuple, Set

try:
    import yaml  # type: ignore
except Exception as e:  # pragma: no cover
    yaml = None

# Optional: DOCX export
try:
    from docx import Document  # type: ignore
    from docx.shared import Pt
except Exception:
    Document = None  # type: ignore

STOPWORDS = set(
    '''a an and are as at be but by for from has have if in into is it its of on or s t that the to with you your we our this those these via using use used than then over under up down out about across after before again further more most other such very can will would should could may might must per'''.split()
)

# ---------------------------- Data Models ----------------------------
@dataclass
class Job:
    id: str
    title: str
    company: str
    location: str
    description: str
    apply_url: str = ""
    source: str = ""

@dataclass
class Bullet:
    text: str
    tags: List[str] = field(default_factory=list)

@dataclass
class Experience:
    role: str
    company: str
    start: str
    end: str
    bullets: List[Bullet]

@dataclass
class Resume:
    name: str
    contact: str
    summary: str
    skills: List[str]
    experiences: List[Experience]
    education: List[str] = field(default_factory=list)
    projects: List[str] = field(default_factory=list)

# ---------------------------- Utilities ----------------------------

def load_jobs_csv(path: Path) -> List[Job]:
    jobs: List[Job] = []
    if not path.exists():
        sys.exit(f"jobs file not found: {path}")
    with path.open(newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        req = {"id", "title", "company", "location", "description"}
        missing = req - set(reader.fieldnames or [])
        if missing:
            sys.exit(f"jobs.csv missing columns: {', '.join(sorted(missing))}")
        for row in reader:
            jobs.append(Job(
                id=str(row.get("id", "")).strip() or str(len(jobs)+1),
                title=row.get("title", "").strip(),
                company=row.get("company", "").strip(),
                location=row.get("location", "").strip(),
                description=row.get("description", "").strip(),
                apply_url=row.get("apply_url", "").strip(),
                source=row.get("source", "").strip(),
            ))
    return jobs


def normalize(s: str) -> str:
    return re.sub(r"[^a-z0-9\s]", " ", s.lower())


def tokenize(text: str) -> List[str]:
    return [w for w in normalize(text).split() if w and w not in STOPWORDS]


def extract_keywords(text: str, top_k: int = 25) -> List[str]:
    from collections import Counter
    counts = Counter(tokenize(text))
    return [w for w, _ in counts.most_common(top_k)]


def match_score(bullet: Bullet, kwset: Set[str]) -> int:
    score = 0
    # tags have higher weight than plain keyword overlaps
    for t in bullet.tags:
        if t.lower() in kwset:
            score += 3
    for tok in tokenize(bullet.text):
        if tok in kwset:
            score += 1
    return score


def load_resume_yaml(path: Path) -> Resume:
    if yaml is None:
        sys.exit("PyYAML is required. Run: pip install pyyaml")
    with path.open("r", encoding="utf-8") as f:
        data = yaml.safe_load(f)
    # Parse bullets with optional tags: "text | tags: python,sql"
    experiences: List[Experience] = []
    for exp in data.get("experiences", []):
        parsed_bullets: List[Bullet] = []
        for raw in exp.get("bullets", []):
            if isinstance(raw, dict):
                parsed_bullets.append(Bullet(text=raw.get("text", ""), tags=raw.get("tags", []) or []))
            else:
                text = str(raw)
                tags: List[str] = []
                if "| tags:" in text:
                    parts = text.split("| tags:")
                    text = parts[0].strip()
                    tags = [t.strip() for t in parts[1].split(',') if t.strip()]
                parsed_bullets.append(Bullet(text=text, tags=tags))
        experiences.append(Experience(
            role=exp.get("role", ""),
            company=exp.get("company", ""),
            start=exp.get("start", ""),
            end=exp.get("end", ""),
            bullets=parsed_bullets,
        ))
    return Resume(
        name=data.get("name", ""),
        contact=data.get("contact", ""),
        summary=data.get("summary", ""),
        skills=data.get("skills", []),
        experiences=experiences,
        education=data.get("education", []),
        projects=data.get("projects", []),
    )


def format_resume_markdown(resume: Resume, top_bullets: List[Tuple[Experience, Bullet]], job: Job) -> str:
    lines: List[str] = []
    lines.append(f"# {resume.name}")
    lines.append(resume.contact)
    lines.append("")
    lines.append(f"**Target role:** {job.title} @ {job.company} – {job.location}")
    lines.append("")
    if resume.summary:
        lines.append("## Summary")
        lines.append(resume.summary)
        lines.append("")
    if resume.skills:
        lines.append("## Skills")
        lines.append(", ".join(resume.skills))
        lines.append("")
    lines.append("## Experience (selected, matched to role)")
    current_section: Optional[str] = None
    for exp, b in top_bullets:
        if current_section != f"{exp.role} — {exp.company}":
            current_section = f"{exp.role} — {exp.company}"
            lines.append(f"**{exp.role}**, {exp.company}  _{exp.start}–{exp.end}_")
        lines.append(f"• {b.text}")
    if resume.education:
        lines.append("")
        lines.append("## Education")
        for edu in resume.education:
            lines.append(f"• {edu}")
    if resume.projects:
        lines.append("")
        lines.append("## Projects")
        for p in resume.projects:
            lines.append(f"• {p}")
    return "\n".join(lines)


def export_docx(markdown_text: str, out_path: Path) -> None:
    if Document is None:
        return
    doc = Document()
    for line in markdown_text.splitlines():
        if line.startswith('# '):
            p = doc.add_heading(level=0)
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(18)
            continue
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        elif line.startswith('**') and '**' in line[2:]:
            # naive bold segment handling
            para = doc.add_paragraph()
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part.strip('*'))
                    run.bold = True
                else:
                    para.add_run(part)
            continue
        else:
            doc.add_paragraph(line)
    doc.save(out_path)


def select_top_bullets(resume: Resume, job: Job, per_experience_limit: int = 3, total_limit: int = 12) -> List[Tuple[Experience, Bullet]]:
    kw = set(extract_keywords(job.description, top_k=60))
    scored: List[Tuple[int, Experience, Bullet]] = []
    for exp in resume.experiences:
        for b in exp.bullets:
            scored.append((match_score(b, kw), exp, b))
    # per‑experience cap first
    per_exp: List[Tuple[Experience, Bullet]] = []
    for exp in resume.experiences:
        these = sorted([(s, e, blt) for (s, e, blt) in scored if e is exp], key=lambda x: x[0], reverse=True)[:per_experience_limit]
        per_exp.extend([(e, blt) for (_, e, blt) in these])
    # then global top, preserving order by score
    global_top = sorted([(match_score(b, kw), exp, b) for exp, b in per_exp], key=lambda x: x[0], reverse=True)
    unique: List[Tuple[Experience, Bullet]] = []
    seen: Set[str] = set()
    for score, exp, b in global_top:
        key = f"{exp.company}|{b.text}"
        if key not in seen and score > 0:
            unique.append((exp, b))
            seen.add(key)
        if len(unique) >= total_limit:
            break
    # if nothing matched, fall back to first bullets
    if not unique:
        for exp in resume.experiences:
            for b in exp.bullets[:2]:
                unique.append((exp, b))
                if len(unique) >= total_limit:
                    break
            if len(unique) >= total_limit:
                break
    return unique


def make_cover_letter(resume: Resume, job: Job, matched_keywords: List[str]) -> str:
    mk = ", ".join(matched_keywords[:10])
    today = datetime.now().strftime("%B %d, %Y")
    return textwrap.dedent(f"""
    {today}

    Hiring Team
    {job.company}

    Re: {job.title}

    Dear Hiring Team,

    I’m excited to apply for the {job.title} role at {job.company}. The description emphasizes {mk}, which align closely with my background. In prior roles, I:\
    – Delivered results that map to the role’s needs (selected bullets appear in the tailored resume).\
    – Collaborated across teams and adapted quickly to new tools and workflows.\
    – Focused on measurable impact while keeping quality and deadlines front‑of‑mind.

    I’d welcome the chance to contribute at {job.company}. Thank you for your time and consideration.

    Sincerely,
    {resume.name}
    {resume.contact}
    """)

# ---------------------------- Commands ----------------------------

def cmd_filter(args: argparse.Namespace) -> None:
    jobs = load_jobs_csv(Path(args.jobs))
    inc = [s.strip().lower() for s in (args.include or '').split(',') if s.strip()]
    exc = [s.strip().lower() for s in (args.exclude or '').split(',') if s.strip()]
    locs = [s.strip().lower() for s in (args.locations or '').split(';') if s.strip()]

    def keep(job: Job) -> bool:
        title = job.title.lower()
        loc = job.location.lower()
        desc = job.description.lower()
        if inc and not any(k in title or k in desc for k in inc):
            return False
        if exc and any(k in title or k in desc for k in exc):
            return False
        if locs and not any(l in loc for l in locs):
            return False
        return True

    kept = [j for j in jobs if keep(j)]
    print(f"Matched {len(kept)} of {len(jobs)} jobs:\n")
    for j in kept:
        print(f"[{j.id}] {j.title} — {j.company} — {j.location}")
        if j.apply_url:
            print(f"  apply: {j.apply_url}")
        if j.source:
            print(f"  source: {j.source}")
        print()


def cmd_tailor(args: argparse.Namespace) -> None:
    jobs = load_jobs_csv(Path(args.jobs))
    target = next((j for j in jobs if str(j.id) == str(args.job_id)), None)
    if not target:
        sys.exit(f"Job id {args.job_id} not found in {args.jobs}")
    resume = load_resume_yaml(Path(args.resume))

    top_bullets = select_top_bullets(resume, target)
    md = format_resume_markdown(resume, top_bullets, target)

    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    md_path = out_dir / f"{target.id}_resume.md"
    md_path.write_text(md, encoding='utf-8')

    if Document is not None:
        docx_path = out_dir / f"{target.id}_resume.docx"
        export_docx(md, docx_path)

    # Cover letter
    matched_kw = extract_keywords(target.description, top_k=20)
    cl = make_cover_letter(resume, target, matched_kw)
    cl_path = out_dir / f"{target.id}_cover_letter.md"
    cl_path.write_text(cl, encoding='utf-8')

    print(f"Wrote {md_path}")
    if Document is not None:
        print(f"Wrote {out_dir / f'{target.id}_resume.docx'}")
    print(f"Wrote {cl_path}")


def cmd_open(args: argparse.Namespace) -> None:
    jobs = load_jobs_csv(Path(args.jobs))
    target = next((j for j in jobs if str(j.id) == str(args.job_id)), None)
    if not target:
        sys.exit(f"Job id {args.job_id} not found in {args.jobs}")
    if not target.apply_url:
        sys.exit("This job has no apply_url in jobs.csv")
    print(f"Opening {target.apply_url}…")
    webbrowser.open(target.apply_url)


# ---------------------------- CLI ----------------------------

def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(prog="job_apply_mvp.py", description="Job filter + resume tailoring (TOS‑friendly)")
    sub = p.add_subparsers(dest="cmd", required=True)

    p.add_argument("--jobs", default="jobs.csv", help="Path to jobs CSV")
    p.add_argument("--resume", default="master_resume.yaml", help="Path to master resume YAML")

    f = sub.add_parser("filter", help="Filter jobs by location/keywords")
    f.add_argument("--locations", default="", help="Semicolon‑separated location substrings, e.g. 'Sacramento, CA; Remote'")
    f.add_argument("--include", default="", help="Comma keywords required in title/description")
    f.add_argument("--exclude", default="", help="Comma keywords to exclude")
    f.set_defaults(func=cmd_filter)

    t = sub.add_parser("tailor", help="Tailor resume & cover letter for one job id")
    t.add_argument("--job-id", required=True, help="ID from jobs.csv")
    t.add_argument("--out-dir", default="out", help="Output directory")
    t.set_defaults(func=cmd_tailor)

    o = sub.add_parser("open", help="Open the job apply URL in your browser")
    o.add_argument("--job-id", required=True, help="ID from jobs.csv")
    o.set_defaults(func=cmd_open)

    return p


def main(argv: Optional[List[str]] = None) -> None:
    args = build_parser().parse_args(argv)
    args.func(args)


if __name__ == "__main__":
    main()

# ---------------------------- SAMPLE FILES ----------------------------
# Copy these samples into your repo as master_resume.yaml and jobs.csv, then edit.

# ==== SAMPLE: master_resume.yaml ====
# name: Wali Basharmal
# contact: West Sacramento, CA • heisenberg@example.com • linkedin.com/in/yourhandle • github.com/yourgit
# summary: >
#   Detail‑oriented student with experience in customer service, logistics, and tutoring.
#   Interested in operations, retail leadership, and data‑driven process improvement.
# skills:
#   - Customer Service
#   - Excel / Google Sheets
#   - Python (basics)
#   - Communication
#   - Teamwork
# experiences:
#   - role: Student Brand Ambassador
#     company: Storage Scholars
#     start: 2024
#     end: 2024
#     bullets:
#       - "Drove sign‑ups by tabling and outreach across campus | tags: sales, outreach, communication"
#       - "Managed logistics tables and inventory handoffs | tags: logistics, operations"
#       - "Collaborated in a small team to meet daily targets | tags: teamwork"
#   - role: Tutor (Volunteer)
#     company: Local Programs
#     start: 2023
#     end: 2024
#     bullets:
#       - "Tutored K‑12 students in math and languages | tags: education, communication"
#       - "Prepared lesson outlines and tracked progress | tags: organization"
# education:
#   - "UC Davis – B.A., International Relations (in progress)"
# projects:
#   - "Built a personal website for a jewelry brand"

# ==== SAMPLE: jobs.csv ====
# id,title,company,location,description,apply_url,source
# 42,Store Executive Intern,Target,West Sacramento, CA,"Work with leaders to manage store operations, drive sales, and deliver guest experience. Skills: communication, leadership, operations, retail, Excel.",https://careers.target.com/internship,Target Careers
# 77,Operations Intern,CVS Health,Sacramento, CA,"Assist store teams with operations, logistics, and reporting. Skills: teamwork, communication, Excel, customer service.",https://jobs.cvshealth.com/internships,CVS Careers
